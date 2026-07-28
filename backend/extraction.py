"""Single multilingual text extractor, replacing the four competing extractors.

Text -> extract(); OCR runs against a hosted model (no local OCR binary), since
Vercel's serverless Python runtime cannot install system packages like tesseract.
"""
from __future__ import annotations

import io
import logging
import os
import tempfile
from dataclasses import dataclass, field
from pathlib import Path

import fitz  # PyMuPDF
import pandas as pd
from docx import Document
from gradio_client import Client, handle_file
from PIL import Image

from languages import detect_script

logger = logging.getLogger(__name__)

OCR_DPI = 300
MIN_TEXT_LEN_PER_PAGE = 20  # below this, a PDF page is treated as a scan and OCR'd

_ocr_client: Client | None = None


def _get_ocr_client() -> Client:
    """Lazily constructed so HF_OCR_SPACE/HF_TOKEN are read after main.py's load_dotenv() has run."""
    global _ocr_client
    if _ocr_client is None:
        space = os.getenv("HF_OCR_SPACE", "baidu/Unlimited-OCR")
        token = os.getenv("HF_TOKEN")  # optional, improves ZeroGPU queue priority
        _ocr_client = Client(space, hf_token=token)
    return _ocr_client


@dataclass
class ExtractedDoc:
    text: str
    pages: list[str] = field(default_factory=list)
    detected_langs: list[str] = field(default_factory=list)
    used_ocr: bool = False


def _ocr_image(img: Image.Image) -> str:
    """Runs OCR via the hosted baidu/Unlimited-OCR Space (ZeroGPU-backed, no local binary needed)."""
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            img.save(tmp, format="PNG")
            tmp_path = tmp.name

        result = _get_ocr_client().predict(
            image_path=handle_file(tmp_path),
            mode="gundam",
            prompt="document parsing.",
            api_name="/run_ocr",
        )
        text = result.get("text", "") if isinstance(result, dict) else str(result)
        return text.strip()
    except Exception as e:
        logger.warning(f"Hosted OCR call failed: {type(e).__name__}: {e}")
        return ""
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)


def _extract_pdf(data: bytes) -> ExtractedDoc:
    pages: list[str] = []
    used_ocr = False
    doc = fitz.open(stream=data, filetype="pdf")
    try:
        for page in doc:
            text = page.get_text().strip()
            if len(text) < MIN_TEXT_LEN_PER_PAGE:
                pix = page.get_pixmap(dpi=OCR_DPI)
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                text = _ocr_image(img).strip()
                used_ocr = True
            pages.append(text)
    finally:
        doc.close()
    return ExtractedDoc(text="\n\n".join(pages), pages=pages, used_ocr=used_ocr)


def _extract_image(data: bytes) -> ExtractedDoc:
    img = Image.open(io.BytesIO(data))
    text = _ocr_image(img).strip()
    return ExtractedDoc(text=text, pages=[text], used_ocr=True)


def _extract_docx(data: bytes) -> ExtractedDoc:
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    text = "\n".join(parts)
    return ExtractedDoc(text=text, pages=[text])


def _extract_tabular(data: bytes, filename: str) -> ExtractedDoc:
    ext = Path(filename).suffix.lower()
    df = pd.read_csv(io.BytesIO(data)) if ext == ".csv" else pd.read_excel(io.BytesIO(data))
    text = df.to_markdown(index=False)
    return ExtractedDoc(text=text, pages=[text])


def extract(file_bytes: bytes, filename: str) -> ExtractedDoc:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        result = _extract_pdf(file_bytes)
    elif ext in (".jpg", ".jpeg", ".png"):
        result = _extract_image(file_bytes)
    elif ext == ".docx":
        result = _extract_docx(file_bytes)
    elif ext in (".xlsx", ".csv"):
        result = _extract_tabular(file_bytes, filename)
    elif ext in (".txt", ".md"):
        text = file_bytes.decode("utf-8", errors="ignore")
        result = ExtractedDoc(text=text, pages=[text])
    else:
        result = ExtractedDoc(text="", pages=[])
    result.detected_langs = detect_script(result.text) if result.text else ["latin"]
    return result
